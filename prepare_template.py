import docx
import os

def prepare_template(file_path, output_path):
    doc = docx.Document(file_path)
    
    replacements = {
        "PRATIWI NOFIANI, S.H., S.I.K., M.M.": "${PIHAK1_NAMA}",
        "KOMPOL": "${PIHAK1_PANGKAT}",
        "88110790": "${PIHAK1_NRP}",
        "PS. Kabag Bekum": "${PIHAK1_JABATAN}",
        "Spripim Polda NTB": "${SATKER_NAMA}",
        "Januari": "${BULAN}",
        "2026": "${TAHUN}",
        "1.413": "${TOTAL_PERTAMAX}",
        "302": "${TOTAL_PERTAMINADEX}",
        "KAYANMA POLDA NTB": "${TTD_JABATAN}",
        "AKBP": "${TTD_PANGKAT}",
        "724707240": "${TTD_NRP}",
        "KABAG YANMA POLDA NTB": "${TTD_JABATAN}", # Backup
    }

    def replace_text_in_paras(paras):
        for para in paras:
            for run in para.runs:
                for old, new in replacements.items():
                    if old in run.text:
                        run.text = run.text.replace(old, new)

    replace_text_in_paras(doc.paragraphs)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_in_paras(cell.paragraphs)

    doc.save(output_path)
    print(f"Template saved to {output_path}")

if __name__ == "__main__":
    template_in = r"e:\spbp\storage\app\templates\BA_Template.docx"
    prepare_template(template_in, template_in)
