import docx
import sys

def extract_text(file_path):
    try:
        doc = docx.Document(file_path)
        print("--- CONTENT ---")
        for para in doc.paragraphs:
            if para.text.strip():
                print(para.text)
        
        print("\n--- TABLES ---")
        for i, table in enumerate(doc.tables):
            print(f"\nTable {i+1}:")
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                print(" | ".join(row_data))
                
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    if len(sys.argv) > 1:
        extract_text(sys.argv[1])
    else:
        print("Usage: python read_docx.py <file_path>")
